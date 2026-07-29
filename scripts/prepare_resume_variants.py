"""Create private, contact-filled PDF-ready resume DOCX variants from the resume pack."""

from __future__ import annotations

import argparse
import io
import zipfile
from pathlib import Path

from docx import Document


TEMPLATES = {
    "01_AI_Agent_Development_CN.docx": "Ronin_AI_Agent_Development_CN.docx",
    "06_AI_Agent_LLM_Intern_EN.docx": "Ronin_AI_Agent_LLM_Intern_EN.docx",
}


def replace_in_document(document: Document, replacements: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            for source, destination in replacements.items():
                run.text = run.text.replace(source, destination)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_document(cell, replacements)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--pack", type=Path, required=True)
    parser.add_argument("--outdir", type=Path, required=True)
    parser.add_argument("--name", required=True)
    parser.add_argument("--phone", required=True)
    parser.add_argument("--email", required=True)
    args = parser.parse_args()

    args.outdir.mkdir(parents=True, exist_ok=True)
    replacements = {
        "Ronin": args.name,
        "[[PHONE]]": args.phone,
        "[[EMAIL]]": args.email,
        "Remote / Nantong / Nanjing / Shanghai / Suzhou / Hangzhou": "Remote / Nantong (Chongchuan) / Nanjing (Pukou)",
        "远程 / 南通 / 南京 / 上海 / 苏州 / 杭州": "远程 / 南通市崇川区 / 南京市浦口区",
        "Remote / Yangtze River Delta": "Remote / Nantong (Chongchuan) / Nanjing (Pukou)",
    }
    safe_name = "".join(character for character in args.name if character.isalnum() or character in "-_ ").strip().replace(" ", "_") or "Candidate"

    with zipfile.ZipFile(args.pack) as archive:
        for source_name, output_name in TEMPLATES.items():
            member = next(name for name in archive.namelist() if name.endswith(source_name))
            document = Document(io.BytesIO(archive.read(member)))
            replace_in_document(document, replacements)
            document.save(args.outdir / output_name.replace("Ronin", safe_name))


if __name__ == "__main__":
    main()
